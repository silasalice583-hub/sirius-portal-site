from __future__ import annotations

import html
import json
from html.parser import HTMLParser
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "exports" / "公众号文章备份-向明与苍焰-2026-08-30.json"
OUTPUT = ROOT / "exports" / "公众号文章备份-向明与苍焰-2026-08-30.docx"

FONT = "Microsoft YaHei"
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DEEP_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x68, 0x72, 0x7D)
LIGHT = RGBColor(0xE8, 0xF0, 0xF7)


class BlockParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.blocks: list[dict[str, str]] = []
        self.current_class = ""
        self.current_text: list[str] | None = None

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        attrs_dict = {key: value or "" for key, value in attrs}
        if tag == "p":
            self.current_class = attrs_dict.get("class", "")
            self.current_text = []
        elif tag == "img":
            src = attrs_dict.get("src", "")
            if src:
                self.blocks.append({"type": "image", "src": src})

    def handle_data(self, data: str) -> None:
        if self.current_text is not None:
            self.current_text.append(data)

    def handle_endtag(self, tag: str) -> None:
        if tag == "p" and self.current_text is not None:
            text = html.unescape("".join(self.current_text)).strip()
            if text:
                self.blocks.append({"type": "paragraph", "class": self.current_class, "text": text})
            self.current_text = None
            self.current_class = ""


def set_run_font(run, size: float, *, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_bottom_border(paragraph, color: str = "D5DFE8", size: str = "8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr_text, fld_end])
    set_run_font(run, 8.5, color=MUTED)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.30)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DEEP_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = doc.styles[name]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(3)
    run = hp.add_run("晶彩未来公众号  ·  文章更新")
    set_run_font(run, 8.5, bold=True, color=MUTED)
    add_bottom_border(hp, "CFD9E3", "6")

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(3)
    run = fp.add_run("公众号文章备份  ·  ")
    set_run_font(run, 8.5, color=MUTED)
    add_page_field(fp)


def add_cover_page(doc: Document, article_count: int) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(116)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("文章更新  ·  完整归档")
    set_run_font(run, 10.5, bold=True, color=BLUE)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(14)
    title.paragraph_format.keep_with_next = True
    run = title.add_run("向明与苍焰\n公众号文章备份")
    set_run_font(run, 28, bold=True, color=DEEP_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    run = subtitle.add_run(f"晶彩未来公众号  ·  {article_count} 篇文章")
    set_run_font(run, 13, color=MUTED)

    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rule.paragraph_format.left_indent = Inches(1.3)
    rule.paragraph_format.right_indent = Inches(1.3)
    rule.paragraph_format.space_after = Pt(24)
    add_bottom_border(rule, "2E74B5", "12")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.line_spacing = 1.45
    run = meta.add_run("导出日期：2026年8月30日\n包含原封面、正文、公众号配图与原文署名\n本地网站分类：文章更新")
    set_run_font(run, 10.5, color=MUTED)


def image_size(path: Path, max_width: float, max_height: float) -> tuple[float, float]:
    with Image.open(path) as image:
        width, height = image.size
    if width <= 0 or height <= 0:
        return max_width, max_height
    scale = min(max_width / width, max_height / height)
    return width * scale, height * scale


def add_image(doc: Document, path: Path, *, cover: bool = False) -> None:
    if not path.exists():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(f"[图片缺失：{path.name}]")
        set_run_font(run, 9.5, color=MUTED)
        return
    max_height = 320 if cover else 470
    width, height = image_size(path, 624, max_height)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(10)
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width / 96), height=Inches(height / 96))


def add_article_paragraph(doc: Document, text: str, css_class: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    size = 11
    color = RGBColor(0x20, 0x20, 0x20)
    bold = False
    if "wechat-quote" in css_class:
        color = MUTED
        paragraph.paragraph_format.left_indent = Inches(0.18)
        paragraph.paragraph_format.right_indent = Inches(0.18)
    elif "wechat-source-line" in css_class:
        size = 10
        color = MUTED
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.space_after = Pt(9)
    elif "wechat-signature" in css_class:
        size = 12
        color = BLUE
        bold = True
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(14)
    run = paragraph.add_run(text)
    set_run_font(run, size, bold=bold, color=color)


def add_article(doc: Document, article: dict[str, object], index: int, total: int) -> None:
    kicker = doc.add_paragraph()
    kicker.paragraph_format.page_break_before = True
    kicker.paragraph_format.space_before = Pt(0)
    kicker.paragraph_format.space_after = Pt(5)
    kicker.paragraph_format.keep_with_next = True
    run = kicker.add_run(f"文章更新  ·  {index:02d}/{total:02d}")
    set_run_font(run, 9.5, bold=True, color=BLUE)

    title = doc.add_paragraph(style="Heading 1")
    title.paragraph_format.space_before = Pt(0)
    run = title.add_run(str(article["title"]))
    set_run_font(run, 16, bold=True, color=BLUE)

    author = str(article.get("author") or article.get("canonicalAuthor") or "")
    date = str(article.get("date") or "")
    metadata = doc.add_paragraph()
    metadata.paragraph_format.space_after = Pt(8)
    metadata.paragraph_format.keep_with_next = True
    run = metadata.add_run(f"作者：{author}    发布日期：{date}    来源：晶彩未来")
    set_run_font(run, 9.5, color=MUTED)
    add_bottom_border(metadata)

    cover = ROOT / str(article["cover"])
    add_image(doc, cover, cover=True)

    source = doc.add_paragraph()
    source.paragraph_format.space_after = Pt(10)
    source.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = source.add_run(f"公众号原文：{article.get('sourceUrl', '')}")
    set_run_font(run, 8.5, color=MUTED)

    parser = BlockParser()
    parser.feed(str(article.get("html") or ""))
    for block in parser.blocks:
        if block["type"] == "image":
            add_image(doc, ROOT / block["src"])
        else:
            add_article_paragraph(doc, block["text"], block.get("class", ""))


def main() -> None:
    payload = json.loads(SOURCE.read_text(encoding="utf-8"))
    articles = list(payload["articles"])
    articles.sort(key=lambda item: (str(item.get("date", "")), str(item.get("title", ""))), reverse=True)

    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "公众号文章备份：向明与苍焰"
    doc.core_properties.subject = "晶彩未来公众号文章更新完整归档"
    doc.core_properties.author = "晶彩未来"
    doc.core_properties.keywords = "公众号, 向明, 苍焰, 文章更新, 备份"

    add_cover_page(doc, len(articles))
    for index, article in enumerate(articles, 1):
        add_article(doc, article, index, len(articles))

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(json.dumps({"output": str(OUTPUT), "articles": len(articles)}, ensure_ascii=False))


if __name__ == "__main__":
    main()
